from __future__ import annotations

from collections.abc import Sequence

from docx import Document as doc_create


class Column:
    def __init__(self, pos: int, total_inputs: int, first_col: Column | None = None, latex: str | None = None, values: Sequence[int] | None = None):
        is_input_variable_column = pos != -1
        if is_input_variable_column:
            if pos > total_inputs or pos == 0:
                raise Exception("Invalid column position")

            construction_values = []
            
            # there are two states: 0 and 1
            # For the first column in a total of three columns
            # there are 8 / 2 = 4 number of each state
            # this means that first column has four zero's and four one's
            number_of_each_state = 2 ** (total_inputs - pos)
            total_entries = 2 ** total_inputs
            
            while len(construction_values) < total_entries:
                for _ in range(number_of_each_state):
                    construction_values.append(0)
                
                for _ in range(number_of_each_state):
                    construction_values.append(1)

            self.__values = tuple(construction_values)
        else:
            self.__values = tuple(values)
        self.__nrows = len(self.__values)
        
        self.latex = latex or chr(ord('A') - 1 + pos)
        
        self.first_col = first_col
        if first_col is not None:
            first_col.table.append([self.latex, *self.values])
        else:
            self.table = [[self.latex, *self.values]]
            self.first_col = self

    @property
    def values(self) -> tuple[int, ...]:
        return self.__values

    @property
    def nrows(self) -> int:
        return self.__nrows
    
    def and_op(self, other: Column) -> Column:
        new_values = tuple(map(lambda vals: int(vals[0] and vals[1]), zip(self.values, other.values, strict=True)))
        new_col_latex = fr"{self.latex} \cdot {other.latex}"
        new_col = Column(-1, -1, latex=new_col_latex, values=new_values, first_col=self.first_col)
        return new_col
    
    def or_op(self, other: Column) -> Column:
        new_values = tuple(map(lambda vals: int(vals[0] or vals[1]), zip(self.values, other.values)))
        new_col_latex = f"{self.latex} + {other.latex}"
        new_col = Column(-1, -1, latex=new_col_latex, values=new_values, first_col=self.first_col)
        return new_col
    
    def not_op(self) -> Column:  
        new_values = tuple(map(lambda x: int(not x), self.values))
        new_col_latex = fr"\overline{{{self.latex}}}"
        new_col = Column(-1, -1, latex=new_col_latex, values=new_values, first_col=self.first_col)
        return new_col

    def __mul__(self, other: Column) -> Column:
        return self.and_op(other)
    
    def __and__(self, other: Column) -> Column:
        return self.and_op(other)
    
    def __or__(self, other: Column) -> Column:
        return self.or_op(other)
    
    def __add__(self, other: Column) -> Column:
        return self.or_op(other)
    
    def __invert__(self) -> Column:
        return self.not_op()


def make_table_document(first_col: Column, filename_to_save: str):
    document = doc_create()
    table = first_col.table

    ncols = len(table)
    nrows = len(table[0])

    # `table` variable contains data for creating
    # the table in Word document.
    # `docx_table` is the table in Word document.
    docx_table = document.add_table(rows=nrows, cols=ncols)
    docx_table.style = "Table Grid"

    for docx_table_row_index in range(nrows):
        current_row_cells = docx_table.rows[docx_table_row_index].cells

        for docx_table_col_index in range(ncols):
            current_cell = current_row_cells[docx_table_col_index]
            current_value = table[docx_table_col_index][docx_table_row_index]

            current_cell      \
            .add_paragraph()  \
            .add_run(
                text=str(current_value)
            )
    
    document.save(filename_to_save)
